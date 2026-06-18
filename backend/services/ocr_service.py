"""
OCR & Document Text Extraction Service
Handles: PDF (text-based), PDF (scanned), images (PNG/JPG), DOCX, plain text
Robust preprocessing for handwritten + printed medical documents.
"""
import io
import logging

logger = logging.getLogger(__name__)


def extract_text(file_obj) -> dict:
    filename = file_obj.filename.lower() if hasattr(file_obj, 'filename') else ""
    content = file_obj.read()
    file_obj.seek(0)

    result = {"text": "", "file_type": "unknown", "page_count": 1,
              "confidence": 0.0, "warnings": []}
    try:
        if filename.endswith('.pdf'):
            result = _extract_pdf(content)
        elif filename.endswith(('.png', '.jpg', '.jpeg', '.tiff', '.bmp', '.webp')):
            result = _extract_image(content)
        elif filename.endswith('.docx'):
            result = _extract_docx(content)
        elif filename.endswith(('.txt', '.text')):
            result = _extract_text_file(content)
        else:
            # Unknown extension — try PDF first, then image
            try:
                result = _extract_pdf(content)
            except Exception:
                result = _extract_image(content)
    except Exception as e:
        import traceback
        traceback.print_exc()
        logger.error(f"Extraction error: {e}")
        result["warnings"].append(str(e))

    result["text"] = _clean_text(result.get("text", ""))
    logger.info(f"Extracted {len(result['text'])} chars, type={result['file_type']}, conf={result['confidence']}")
    return result


def _preprocess(img):
    """
    Safe, robust preprocessing for medical docs.
    Works on printed forms, handwriting, photocopies, scans.
    """
    from PIL import Image, ImageFilter, ImageEnhance, ImageOps

    # Convert to grayscale
    img = img.convert("L")

    # Upscale if small — Tesseract accuracy degrades under ~200 DPI
    w, h = img.size
    if w < 1200:
        scale = max(2, 1200 // w)
        img = img.resize((w * scale, h * scale), Image.LANCZOS)

    # Boost contrast (helps faded ink, photocopies)
    img = ImageEnhance.Contrast(img).enhance(2.0)

    # Sharpen
    img = img.filter(ImageFilter.SHARPEN)

    # Auto-level histogram (handles uneven exposure)
    try:
        img = ImageOps.autocontrast(img, cutoff=2)
    except Exception:
        pass

    return img


def _tesseract_extract(img) -> tuple:
    """
    Run Tesseract with multiple PSM modes, return (best_text, avg_confidence).
    PSM 4 = single column (good for narrow prescriptions)
    PSM 6 = uniform block (good for structured forms)
    PSM 11 = sparse text (good for scattered handwriting)
    PSM 3 = auto (default fallback)
    """
    import pytesseract

    configs = [
        "--psm 4 --oem 1",
        "--psm 6 --oem 1",
        "--psm 11 --oem 1",
        "--psm 3 --oem 1",
    ]

    best_text = ""
    for cfg in configs:
        try:
            t = pytesseract.image_to_string(img, config=cfg, lang="eng")
            if len(t.strip()) > len(best_text.strip()):
                best_text = t
        except Exception:
            continue

    # Estimate confidence from best config
    confidence = 0.4
    try:
        data = pytesseract.image_to_data(img, config="--psm 6 --oem 1",
                                          output_type=pytesseract.Output.DICT)
        confs = [int(c) for c in data["conf"] if str(c).lstrip('-').isdigit() and int(c) > 10]
        if confs:
            confidence = sum(confs) / len(confs) / 100
    except Exception:
        pass

    return best_text, round(confidence, 2)


def _extract_image(content: bytes) -> dict:
    from PIL import Image
    img = Image.open(io.BytesIO(content))
    img = _preprocess(img)
    text, confidence = _tesseract_extract(img)

    warnings = []
    if confidence < 0.5:
        warnings.append("Low OCR confidence — handwritten or low-quality scan. Verify before coding.")

    return {"text": text, "file_type": "image", "page_count": 1,
            "confidence": confidence, "warnings": warnings}


def _extract_pdf(content: bytes) -> dict:
    import pdfplumber
    text_parts = []
    page_count = 0

    with pdfplumber.open(io.BytesIO(content)) as pdf:
        page_count = len(pdf.pages)
        for page in pdf.pages:
            t = page.extract_text()
            if t:
                text_parts.append(t)

    combined = "\n\n".join(text_parts)

    # If very little text, treat as scanned PDF and OCR it
    if len(combined.strip()) < 80:
        try:
            return _ocr_pdf(content)
        except Exception as e:
            logger.warning(f"PDF OCR failed: {e}")
            return {"text": combined, "file_type": "pdf", "page_count": page_count,
                    "confidence": 0.3,
                    "warnings": [f"OCR failed: {e}. Install poppler: apt install poppler-utils or brew install poppler"]}

    return {"text": combined, "file_type": "pdf", "page_count": page_count,
            "confidence": 1.0, "warnings": []}


def _ocr_pdf(content: bytes) -> dict:
    from pdf2image import convert_from_bytes
    pages = convert_from_bytes(content, dpi=300)
    texts = []
    confidences = []
    for page_img in pages:
        img = _preprocess(page_img)
        t, c = _tesseract_extract(img)
        texts.append(t)
        confidences.append(c)

    avg_conf = sum(confidences) / len(confidences) if confidences else 0.4
    return {"text": "\n\n".join(texts), "file_type": "pdf_scanned",
            "page_count": len(pages), "confidence": round(avg_conf, 2),
            "warnings": ["OCR used — verify extracted text"] if avg_conf < 0.7 else []}


def _extract_docx(content: bytes) -> dict:
    from docx import Document
    doc = Document(io.BytesIO(content))
    parts = [p.text for p in doc.paragraphs if p.text.strip()]
    for table in doc.tables:
        for row in table.rows:
            r = " | ".join(c.text.strip() for c in row.cells if c.text.strip())
            if r:
                parts.append(r)
    return {"text": "\n".join(parts), "file_type": "docx",
            "page_count": 1, "confidence": 1.0, "warnings": []}


def _extract_text_file(content: bytes) -> dict:
    try:
        text = content.decode("utf-8")
    except UnicodeDecodeError:
        text = content.decode("latin-1")
    return {"text": text, "file_type": "txt", "page_count": 1,
            "confidence": 1.0, "warnings": []}


def _clean_text(text: str) -> str:
    import re
    text = re.sub(r'\n{3,}', '\n\n', text)
    text = re.sub(r' {3,}', ' ', text)
    text = re.sub(r'[^\x20-\x7E\n\r\t°±µ/]', ' ', text)
    return text.strip()
